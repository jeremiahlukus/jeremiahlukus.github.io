from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTE = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x8A, 0x5A, 0x1E)
BODY = "Calibri"

doc = Document()
s = doc.sections[0]
s.top_margin = s.bottom_margin = Inches(0.5)
s.left_margin = s.right_margin = Inches(0.65)

st = doc.styles["Normal"]
st.font.name = BODY
st.font.size = Pt(9.5)
st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.line_spacing = 1.06

def rule(p, size=6, color="C8C8C8"):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "2"); bot.set(qn("w:color"), color)
    b.append(bot); pPr.append(b)

def para(space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def run(p, text, size=9.5, bold=False, italic=False, color=INK, caps=False, spacing=None):
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color; r.font.name = BODY
    if caps: r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(spacing)); rPr.append(sp)
    return r

def heading(text):
    p = para(space_before=11, space_after=5)
    run(p, text, size=9, bold=True, color=ACCENT, caps=True, spacing=30)
    rule(p)

def role(title, org, dates, where=None):
    p = para(space_before=8, space_after=1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)
    run(p, title, size=10.5, bold=True)
    run(p, "  ·  " + org, size=10.5, color=INK)
    run(p, "\t" + dates, size=9, color=MUTE)
    if where:
        q = para(space_after=3)
        run(q, where, size=8.5, italic=True, color=MUTE)

def bullet(text, bold_lead=None, indent=0.0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18 + indent)
    p.paragraph_format.first_line_indent = Inches(-0.13)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.06
    if bold_lead:
        run(p, bold_lead, size=9.5, bold=True)
    run(p, text, size=9.5)
    return p

def kv(label, text):
    p = para(space_after=2)
    run(p, label + "  ", size=9, bold=True, color=MUTE)
    run(p, text, size=9)

# ── Header ───────────────────────────────────────────────────────────────────
p = para()
run(p, "Jeremiah Parrack", size=23, bold=True, spacing=-8)
p = para(space_after=2)
run(p, "Lead Software Engineer  ·  Site Reliability / Platform", size=10, color=ACCENT)
p = para(space_after=1)
run(p, "Atlanta, GA  ·  jeremiahlukus1@gmail.com  ·  github.com/jeremiahlukus  ·  jeremiahlukus.github.io  ·  linkedin.com/in/jeremiahlukus",
    size=8.5, color=MUTE)
rule(para(space_before=3), size=10, color="8A5A1E")

# ── Summary ──────────────────────────────────────────────────────────────────
heading("Summary")
p = para(space_after=2)
run(p, "Software engineer with nine years building and operating production systems on AWS. I work at the seam "
       "between application code and the platform it runs on — shipping Rails, TypeScript, and Flutter applications, "
       "then automating the infrastructure, pipelines, and patching that keep them alive. Currently leading "
       "reliability and platform work across 311 repositories and 26 GitHub organizations at Cox Automotive, "
       "and contributing to the Ruby and Rust open source that runs Rails on AWS Lambda. I publish and maintain "
       "Flutter packages and have shipped four apps of my own to the App Store and Google Play.", size=9.5)

# ── Experience ───────────────────────────────────────────────────────────────
heading("Experience")

role("Lead Software Engineer / SRE", "Manheim — Cox Automotive", "May 2018 – Present",
     "World's largest wholesale auto auction: 145 auction sites across North America, Europe, Asia, and Australia.")
bullet("4,000+ commits across 311 repositories in 26 GitHub organizations.", bold_lead="Fleet-scale ownership: ")
bullet("wrote Python tooling that scans a GitHub Enterprise org for outdated AWS service versions — Elastic Beanstalk "
       "solution stacks, AWS Glue versions, and Terraform releases — opens pull requests with the upgrade, flags "
       "severely outdated versions, and posts Slack notifications on a recurring schedule.",
       bold_lead="Automated dependency patching: ")
bullet("migrated 100+ repositories onto shared, standardized GitHub Actions runners and upgraded reusable Terraform "
       "workflows org-wide, retiring per-team Jenkins pipelines.", bold_lead="CI standardization: ")
bullet("built the Terraform and AWS Glue infrastructure behind DealShield's Assurance Protect Management System — ETL "
       "pipelines for fraud modeling, pending-reason analysis, problem-VIN detection, and inventory — plus Snowflake "
       "production and lower environments.", bold_lead="Data platform: ")
bullet("Terraform for Seller Dashboard and API gateway infrastructure, plus sale-history and disclosure services.",
       bold_lead="Seller Tools: ")
bullet("authored reusable pipeline templates for serverless apps, Angular modules, and npm packages, cutting pipeline "
       "definitions from hundreds of lines to roughly twenty so teams could own their own deploys.",
       bold_lead="Developer experience: ")
bullet("implemented logging, custom metrics, and dashboards that cut hours off production debugging; standardized "
       "alerting and secrets management on AWS Parameter Store across teams.", bold_lead="Observability: ")
bullet("onboarded workloads to the internal service catalog and authored Production Readiness Review documentation.",
       bold_lead="Governance: ")
bullet("ran recurring talks and tutorials on Terraform, deployment ownership, and the release process; mentored "
       "engineers out of the apprenticeship program.", bold_lead="Enablement: ")

role("Founder / Lead Software Engineer", "Jeremiah Parrack LLC", "Feb 2021 – Present",
     "Independent consultancy: client engineering plus my own published apps.")
bullet("full-stack and platform engineering on a project-portfolio-management SaaS — a Rails API with a "
       "React/TypeScript frontend running on AWS Lambda. 215+ commits.",
       bold_lead="Acuity PPM (ongoing): ")
bullet("serverless deploy automation, dedicated per-client production environments, and ephemeral review environments.",
       indent=0.18)
bullet("Lambda access to Postgres on RDS inside a VPC, Hasura, and SQS dead-letter queues with job retry.", indent=0.18)
bullet("New Relic, Sentry, Lambda Insights, log retention, and frontend caching.", indent=0.18)
bullet("secure environment loading through Crypteia and Rails-on-Lambda through Lamby — both projects I contribute "
       "to upstream, so production experience feeds the libraries directly.", indent=0.18)
bullet("four applications shipped to the App Store and Google Play:", bold_lead="Own products: ")
bullet("jiu-jitsu training tracker — offline-first local database with cloud sync. Over 1,000 users. Flutter, Firebase; iOS and Android.",
       bold_lead="FlowJitsu — ", indent=0.18)
bullet("guitar tab and chord library for worship musicians, 6,000+ songs and over 600 users. Flutter, Firebase; iOS.",
       bold_lead="Joyful Noise Tabs — ", indent=0.18)
bullet("AI-narrated reading library that underlines each sentence as it is spoken. Flutter; iOS.",
       bold_lead="Pneuma Stream — ", indent=0.18)
bullet("published to the App Store. Flutter.", bold_lead="Smart Sprout — ", indent=0.18)
bullet("led a small delivery team, translated client requirements into scoped work, and built in-app subscriptions, "
       "release automation, and reusable logging and monitoring packages.", bold_lead="Delivery: ")

role("Software Engineer", "STORD", "Sept 2017 – May 2018",
     "Warehousing and distribution network with custom software for multi-warehouse product flow.")
bullet("built SOAP and REST integrations against third-party warehouse management systems to push and pull inventory data.")
bullet("maintained the Ruby codebase through upgrades, patches, releases, and database migrations; played a key role "
       "in a database redesign.")
bullet("introduced continuous integration, automated error notification for the development team, and Elasticsearch; "
       "oversaw unit and functional testing, debugging, security, and documentation.")

# ── Open source ──────────────────────────────────────────────────────────────
heading("Open Source")

p = para(space_before=1, space_after=3)
run(p, "Ruby & Rust — running Rails on AWS Lambda", size=9, bold=True, color=MUTE, caps=True, spacing=16)
bullet("Rails on AWS Lambda. Added Rack 3 support, carried the gem from Rails 7.1 to 8.1 and onto Ruby 3.4 / "
       "Bundler 4, and made shutdown exit gracefully on SIGTERM.", bold_lead="Lamby — 625★:  ")
bullet("Rust Lambda extension that preloads secure environment variables. Remediated GHSA-82j2-j2ch-gfr8 (a "
       "denial-of-service in rustls-webpki), upgraded the Rust toolchain, and fixed multi-architecture ARM releases.",
       bold_lead="Crypteia — 77★:  ")
bullet("ActiveJob on SQS and Lambda. Modernized CI, added a release workflow, and moved the gem to Ruby 3.4.",
       bold_lead="Lambdakiq — 198★:  ")
bullet("added a Tailwind install task, upgraded RuboCop and fixed the resulting offenses, and repaired the test suite "
       "in the Ruby on Jets asset pipeline.", bold_lead="jetpacker:  ")

p = para(space_before=6, space_after=3)
run(p, "Flutter & Dart", size=9, bold=True, color=MUTE, caps=True, spacing=16)
bullet("Flutter plugin I wrote and maintain on pub.dev, wrapping the native Android and iOS Authorize.Net payment "
       "SDKs behind one Dart API. Ten releases, still shipping.", bold_lead="authorize_net_plugin — pub.dev:  ")
bullet("production Flutter starter — Riverpod state management, authentication flow, structured logging, and a test harness.",
       bold_lead="flutter_template — 27★:  ")
bullet("removed a spurious iOS simulator warning from the community video player used across the Flutter "
       "ecosystem.", bold_lead="chewie — 2.1k★:  ")
bullet("added chorus-style passthrough to the line processor and support for comment directives in this chord "
       "parser and renderer.", bold_lead="flutter_chord:  ")

# ── Skills ───────────────────────────────────────────────────────────────────
heading("Technical Skills")
kv("Languages", "Ruby, Dart, TypeScript, JavaScript, Python, Rust, SQL, Bash")
kv("Frameworks", "Rails, Flutter, React, Ruby on Jets, Hasura, Riverpod")
kv("AWS", "Lambda, Glue, RDS, ECS, Elastic Beanstalk, SQS, SSM Parameter Store, VPC, CloudWatch, S3, IAM")
kv("Infrastructure", "Terraform, AWS CDK, AWS SAM, Docker, Snowflake, Firebase")
kv("CI/CD & Ops", "GitHub Actions, Jenkins, New Relic, Sentry, Lambda Insights, PostgreSQL, Redis, Elasticsearch")

# ── Education ────────────────────────────────────────────────────────────────
heading("Education")
p = para()
p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)
run(p, "Georgia State University", size=10, bold=True)
run(p, "  ·  B.S. Computer Science, concentration in Computer Software Systems", size=9.5)
run(p, "\tMay 2018", size=9, color=MUTE)

out = "JeremiahParrackResume-2026.docx"
doc.save(out)
print("wrote", out)
